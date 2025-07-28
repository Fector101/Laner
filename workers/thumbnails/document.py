from concurrent.futures import ThreadPoolExecutor
from PIL import Image
from abc import ABC, abstractmethod
import os,threading,io,traceback,time
import fitz  # PyMuPDF


if __name__=='image' or __name__=='__main__':
    from base import BaseGen,BaseGenException,use_pc_for_static,DocExtractorABS
    from testing.sword import NetworkManager,NetworkConfig
    from testing.helper import getAppFolder, urlSafePath,_joinPath, getFileExtension,removeFirstDot
else:
    from workers.thumbnails.base import BaseGen,BaseGenException,DocExtractorABS
    from workers.sword import NetworkConfig
    from workers.thumbnails.base import use_pc_for_static
    from workers.helper import gen_unique_filname, urlSafePath, _joinPath, getFileExtension,removeFirstDot


__server_ip = None
__server_port = None



class DocumentIconExtractor(BaseGen):
    # list_of_collected_docs_tuples:list[tuple[str,str]] = [] #(input_path, thumbnail_path)
    ordered_object: dict[str,list[tuple[str,str]]]= {} #{'pdf':[(actual_path,thumbnail_path),(actual_path,thumbnail_path)],'docx':[(actual_path,thumbnail_path),(actual_path,thumbnail_path),...]}

    def __init__(self, doc_path:str='',max_threads=4, server_ip = NetworkConfig.server_ip, server_port = NetworkConfig.port,_thread=True):
        """
        A document path to get thumbnail_url (the only purpose of this arg is to get thumbnail_url)  
        The real work is done by `list_of_collected_docs_tuples` that is share btw instances and clear after `extract` method  
        when i set a `document_path` str for `documents_collection`
        """
        super().__init__(server_ip, server_port)
        global __server_ip,__server_port
        __server_ip = server_ip
        __server_port = server_port
        # print('entered docs extractor:',doc_path)

        self.thumbnail_folder = '.docs-covers'
        self.max_threads = max_threads
        self._thread = _thread

        if doc_path:
            self.document_path = doc_path # Set item_path to path if only trying to gen thumbnail_url
            self._extension = None
            self.extension = self.document_path
        
    @property
    def item_path(self):
        return self.document_path
    
    @property
    def documents_collection(self)-> dict[str,list[tuple[str,str]]]:
        """Getter method"""
        return self.ordered_object
    
    @documents_collection.setter
    def documents_collection(self, document_path:str):
        """Setter method with pre-processing logic"""
        # print("Running logic before setting value")
        
        # Example validation/pre-processing
        if self.extension not in extractors.keys(): # Important so i remember to store all class in `extractors` object
            print(f'Document Thumbnailer not yet avaliable: {self.extension}')
            return
        
        if self.extension not in self.ordered_object:
            self.ordered_object[self.extension]=[]

        self.document_path=document_path # Set item_path to path if only trying to gen thumbnail_url
        self.ordered_object[self.extension].append((document_path,self.thumbnail_path))

        # self.document_path=document_path # Set item_path to path if only trying to gen thumbnail_url
        # self.list_of_collected_docs_tuples.append((document_path,self.thumbnail_path))

    @property
    def extension(self):
        return self._extension
    
    @extension.setter
    def extension(self,document_path):
        if document_path == '':
            self._extension=''
            return
        self._extension = removeFirstDot(getFileExtension(document_path).lower())
    
    def extract(self):
        if self._thread:
            if not isinstance(self.documents_collection,object):
                print(f'Bad input in DocIconExtractor: {self.documents_collection}, excepted object of format: [tuples] of str')
                return

            if self.documents_collection.values():
                threading.Thread(
                    target=self.__assignThreads,
                    daemon=True
                ).start()
        else:
            container = self.documents_collection
            self.ordered_object = {} # Clears self.documents_collection, Don't change this its using a setter referencing `self.ordered_object`
            for extension, list_of_tuples in container.items():
                class__ = extractors[extension]
                if issubclass(class__,DocExtractorABS):
                    class__().extract(*list_of_tuples)
    
    def __assignThreads(self):
        """
        Generate thumbnails from multiple videos in parallel.
        """
        container = self.documents_collection
        self.ordered_object = {} # Clears self.documents_collection, Don't change this its using a setter referencing `self.ordered_object`
        try:
            for extension, list_of_tuples in container.items():
                try:
                    with ThreadPoolExecutor(self.max_threads) as executor:
                        # args_list = [("a", 1), ("b", 2), ("c", 3)]  # List of tuples (arg1, arg2)
                        # executor.map(lambda args: my_func(*args), args_list)
                        # print('documents_collection: ',list_of_tuples)
                        results= list(executor.map(lambda args: extractors[extension]().extract(*args), list_of_tuples))
                        # print("Thread completed processing", results) # This Actually returns list of values
                except Exception as e:
                    print(f'\nDocumentIconExtractor Failed in badge for {extension}: {e}')
                    traceback.print_exc()
                # self.ordered_object[extension] = []
                # print('Done getting thumbnails for:',extension)
        except Exception as e:
            print(f"ThreadPool error: {e}")
            traceback.print_exc()
            # self.ordered_object={}
        finally:
            pass
            # self.list_of_collected_docs_tuples=[] # this will clear in coming stack so bad
            # print("Thread cleanup complete")
    @property
    def pc_static_img(self):
        if self.extension not in extractors.keys():
            return "assets/icons/file.png"
        return use_pc_for_static(f"{self.extension}.png")
    


class PdfExtractor(DocExtractorABS):
    def __init__(self):
        super().__init__()
        self._thumbnail_path = ''
    
    @property
    def thumbnail_path(self):
        return self._thumbnail_path
    
    @property
    def default_img_name(self):
        return removeFirstDot(getFileExtension(self.thumbnail_path).lower())
    
    def extract(self,absolute_filepath,thumbnail_path,zoom=1.0):

        """
        Save the first page of a PDF as an image file
        
        Args:
            absolute_filepath: Complete Path to the input PDF file
            thumbnail_path: Complete Path to save the output image
            zoom: Zoom factor for higher resolution (default 2.0)
        """
        self._thumbnail_path = thumbnail_path
        try:
            # Open the PDF
            doc = fitz.open(absolute_filepath)
            
            # Get the first page
            page = doc.load_page(0)
            
            # Create a matrix for zooming (higher resolution)
            mat = fitz.Matrix(zoom, zoom)
            
            # Get the pixmap (image) of the page
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image
            img = Image.open(io.BytesIO(pix.tobytes()))
            
            # Save the image
            img.save(self.thumbnail_path)
            # print(f"First page saved as {self.thumbnail_path}")
            return 'Sam' #Can return value
        except Exception as e:
            print(f"\nUnexpected PDFExtractor Error---> {e}\nPDF File: {absolute_filepath}\n") 
            traceback.print_exc()
        
        # happen when extraction not saved
        self.getfailSafeImg()


from docx import Document
from PIL import Image
import io
import traceback

class DocxExtractor(DocExtractorABS):
    def __init__(self):
        super().__init__()
        self._thumbnail_path = ''
        self._font_path = self._get_system_font()

    def _get_system_font(self):
        """Get the best available font path"""
        # Try Windows first
        if os.name == 'nt':
            font_path = get_windows_font_path()
            if font_path:
                return font_path
        
        # Try Linux/Mac locations if not found on Windows
        font_paths = [
            '/usr/share/fonts/truetype/msttcorefonts/arial.ttf',
            '/Library/Fonts/Arial.ttf',
            os.path.join(os.path.dirname(__file__), 'fonts/arial.ttf')
        ]
        
        for path in font_paths:
            if os.path.exists(path):
                return path
        
        bundled_font = os.path.join(getAppFolder(), 'fonts', 'LiberationSans-Regular.ttf') # LiberationSans-Regular.ttf is free
        if os.path.exists(bundled_font):
            return bundled_font

        return None  # Will use PIL's basic font
    
    @property
    def thumbnail_path(self):
        return self._thumbnail_path
    
    @property
    def default_img_name(self):
        return removeFirstDot(getFileExtension(self.thumbnail_path).lower())

    def extract(self, absolute_filepath, thumbnail_path, zoom=3.0):
        """
        Save a thumbnail representation of a DOCX file as an image
        
        Args:
            absolute_filepath: Complete path to the input DOCX file
            thumbnail_path: Complete path to save the output image
            zoom: Zoom factor for higher resolution (default 2.0)
        """
        self._thumbnail_path = thumbnail_path
        # print('Received Path:', absolute_filepath)
        
        try:
            # Open the DOCX file
            doc = Document(absolute_filepath)
            
            # Create a blank image (since DOCX doesn't have direct page images like PDF)
            # We'll create a representation of the first page's text
            img = Image.new('RGB', (800, 600), color=(255, 255, 255))
            
            # For demonstration - we'll just show the first paragraph
            # In a real implementation, you might want to render text properly
            first_paragraph = doc.paragraphs[0].text if len(doc.paragraphs) > 0 else "DOCX Document"
            
            # Add text to image (simple representation)
            from PIL import ImageDraw, ImageFont
            draw = ImageDraw.Draw(img)
            font_size = int(40 * zoom)
            try:
                if self._font_path:
                    font = ImageFont.truetype(self._font_path, font_size)
                    # print(f'Using {self._font_path}')
                    # font = ImageFont.truetype("arial.ttf", 40)
                else:
                    font = ImageFont.load_default()
                    print("Could Find Font Using basic default font - install Arial for better results")
            except:
                font = ImageFont.load_default()
            
            
            # text = doc.paragraphs[0].text if len(doc.paragraphs) > 0 else "Word Document"
            # draw.text((50, 50), text, fill=(0, 0, 0), font=font)
            draw.text((50, 50), first_paragraph, fill=(0, 0, 0), font=font)
            
            # Save the image
            img.save(self.thumbnail_path)
            # print(f"DOCX thumbnail saved as {self.thumbnail_path}")
            return 'Success'  # Can return value
            
        except Exception as e:
            print(f"\nUnexpected DocxExtractor Error---> {e}\nDOCX File: {absolute_filepath}\n") 
            traceback.print_exc()
        self.getfailSafeImg()
        return 'failed'

def get_windows_font_path(font_name="arial.ttf"):
    """Get the full path to a font file on Windows"""
    # Get the Windows directory from environment variables
    windir = os.environ.get('WINDIR', 'C:\\Windows')
    font_path = os.path.join(windir, 'Fonts', font_name)
    
    if os.path.exists(font_path):
        return font_path
    
    # Check alternative locations (some systems may differ)
    alt_paths = [
        os.path.join(os.environ.get('SystemRoot', 'C:\\Windows'), 'Fonts', font_name),
        'C:\\Windows\\Fonts\\' + font_name,
        'D:\\Windows\\Fonts\\' + font_name  # For rare multi-drive installations
    ]
    
    for path in alt_paths:
        if os.path.exists(path):
            return path
    
    return None

from typing import TypedDict, Type
class ExtractorsDict(TypedDict):
    pdf: Type[PdfExtractor]
    docx: Type[DocxExtractor]
    # Add other supported extensions here

extractors: ExtractorsDict = {
    'pdf': PdfExtractor,
    'docx': DocxExtractor
}
extractors['docx']().extract

if __name__ == '__main__':
    server_ip = NetworkManager().get_server_ip()  # Replace with actual server IP if needed
    doc_path=r'c:\Users\hp\Desktop\Linux\my_code\Laner\workers\thumbnails\test.pdf'
    r=DocumentIconExtractor(doc_path,server_ip,_thread=False)
    # print(r.thumbnail_url)



# import sys
# def dump_threads():
#     for thread_id, frame in sys._current_frames().items():
#         print(f"\nThread {thread_id}")
#         traceback.print_stack(frame)
        
# Call dump_threads() when you detect a freeze