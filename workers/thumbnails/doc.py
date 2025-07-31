# from typing import LiteralString
# -> LiteralString | str | bytes
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from PIL import Image
import os, threading, io, traceback, time, psutil, queue
import fitz  # PyMuPDF
from docx import Document
from PIL import ImageDraw, ImageFont
from typing import Dict, List, Optional, Type

import atexit

if __name__ in ['thumbnails.doc']:
    from .base import BaseGen,DocExtractorABS
    from .testing.sword import NetworkManager,NetworkConfig
    from .base import use_pc_for_static
    from .testing.helper import getFileExtension,removeFirstDot,getAppFolder
    NetworkConfig.server_ip=NetworkManager().get_server_ip() # NetworkConfig has an import mismatch from different file when running without GUI and straight from server file
else:
    from workers.thumbnails.base import BaseGen,DocExtractorABS
    from workers.sword import NetworkConfig
    from workers.thumbnails.base import use_pc_for_static
    from workers.helper import getFileExtension,removeFirstDot,getAppFolder


# Import base classes
# (Same import logic as before)
def get_windows_font_path(font_name="arial.ttf"):
    """Get the full path to a font file on Windows"""
    # Get the Windows directory from environment variables
    windir = os.environ.get('WINDIR', 'C:\\Windows')
    font_path = os.path.join(windir, 'Fonts', font_name)

    if os.path.exists(font_path):
        return font_path

    # Check alternative locations (some systems may differ)
    alt_paths = [
        os.path.join(os.environ.get('SystemRoot', 'C:\\Windows'), 'Fonts', font_name),
        'C:\\Windows\\Fonts\\' + font_name,
        'D:\\Windows\\Fonts\\' + font_name  # For rare multi-drive installations
    ]

    for path in alt_paths:
        if os.path.exists(path):
            return path

    return None


def get_system_font():
    """Get the best available font path"""
    # Try Windows first
    if os.name == 'nt':
        font_path = get_windows_font_path()
        if font_path:
            return font_path

    # Try Linux/Mac locations if not found on Windows
    font_paths = [
        '/usr/share/fonts/truetype/msttcorefonts/arial.ttf',
        '/Library/Fonts/Arial.ttf',
        os.path.join(os.path.dirname(__file__), 'fonts/arial.ttf')
    ]

    for path in font_paths:
        if os.path.exists(path):
            return path

    bundled_font = os.path.join(getAppFolder(), 'fonts',
                                'LiberationSans-Regular.ttf')  # LiberationSans-Regular.ttf is free
    if os.path.exists(bundled_font):
        return bundled_font

    return None  # Will use PIL's basic font


# =====================
# Resource-Aware Task Queue
# =====================
class ThumbnailTaskQueue:
    """Centralized queue system with resource monitoring"""
    _instance = None
    _lock = threading.Lock()

    def __new__(cls):
        with cls._lock:
            if cls._instance is None:
                cls._instance = super().__new__(cls)
                cls._instance._init_queue()
            return cls._instance

    def _init_queue(self):
        self.task_queue = queue.Queue()
        self.resource_monitor = ResourceMonitor()
        self.config = ThumbnailConfig()
        self.executor = None
        self.is_running = False
        self.processing_thread = None
        self.extractor_registry = ExtractorRegistry()
        atexit.register(self.shutdown)

    def add_task(self, document_path: str,thumbnail_path:str):
        """Add a document to the processing queue"""
        ext = self._get_extension(document_path)

        if not self.extractor_registry.get_extractor(ext):
            print(f'Thumbnailer not available: {ext}')
            return
        self.task_queue.put((ext, document_path, thumbnail_path))
        self._ensure_processing()

    def _ensure_processing(self):
        """Start processing thread if not already running"""
        if not self.is_running and not self.processing_thread:
            print('_ensure_processing')
            self.is_running = True
            self.processing_thread = threading.Thread(
                target=self._process_queue,
                daemon=True
            )
            self.processing_thread.start()

    def _process_queue(self):
        """Process tasks from the queue with resource awareness"""
        try:
            while self.is_running:
                # Small sleep to prevent tight looping when queue is empty
                time.sleep(0.1)

                # Only check resources when we actually have work to do
                if not self.task_queue.empty():
                    # Check system resources before processing
                    if self.resource_monitor.is_system_overloaded():
                        time.sleep(5)
                        continue

                    # Process batch of tasks
                    batch = self._get_next_batch()
                    print('This is batch:',len(batch))
                    if not batch:
                        continue

                    # Group by extension for efficient processing
                    grouped = self._group_by_extension(batch)

                    # Process each group
                    # print('grouped',grouped)
                    for ext, tasks in grouped.items():
                        self._process_task_group(ext, tasks)
        except Exception as e:
            print(f"Queue processing failed: {e}")
            traceback.print_exc()
        finally:
            self.is_running = False
            self.processing_thread = None

    def _get_next_batch(self) -> list:
        """Get next batch of tasks respecting batch size"""
        batch = []
        while len(batch) < self.config.batch_size and not self.task_queue.empty():
            try:
                batch.append(self.task_queue.get_nowait())
            except queue.Empty:
                break
        return batch

    def _group_by_extension(self, batch: list) -> dict:
        """Group tasks by file extension"""
        grouped = {}
        for ext, doc_path, thumb_path in batch:
            if ext not in grouped:
                grouped[ext] = []
            grouped[ext].append((doc_path, thumb_path))
        return grouped

    def _process_task_group(self, ext: str, tasks: list):
        """Process a group of tasks for the same extension"""
        extractor_class = self.extractor_registry.get_extractor(ext)
        if not extractor_class:
            return

        try:
            # Create executor based on task type
            if ext == 'pdf' and os.name == 'nt':  # Windows
                # On Windows, use threads for PDF processing
                executor_class = ThreadPoolExecutor
                max_workers = min(len(tasks), self.config.max_threads)
            elif self._should_use_multiprocessing(ext):
                executor_class = ProcessPoolExecutor
                max_workers = min(len(tasks), os.cpu_count() or 1)
            else:
                executor_class = ThreadPoolExecutor
                max_workers = min(len(tasks), self.config.max_threads)

            with executor_class(max_workers) as executor:
                # Process all tasks in this group
                futures = []
                for doc_path, thumb_path in tasks:
                    # print('_process_task_group')
                    # future = executor.submit(
                    #     extractor_class().extract,
                    #     doc_path,
                    #     thumb_path,
                    #     self.config
                    # )
                    future = executor.submit(
                        self._safe_extract,
                        extractor_class,
                        doc_path,
                        thumb_path
                    )
                    futures.append(future)

                # Wait for completion but don't block queue processing
                for future in futures:
                    try:
                        future.result(timeout=self.config.task_timeout)
                    except Exception as e:
                        print(f"Task failed: {e}")
        except Exception as e:
            print(f"Failed to process {ext} tasks: {e}")

    def _safe_extract(self, extractor_class, doc_path, thumb_path):
        """Wrapper for safe extraction"""
        try:
            return extractor_class().extract(doc_path, thumb_path, self.config)
        except Exception as e:
            print(f"Extraction failed for {doc_path}: {e}")
            return False

    def _should_use_multiprocessing(self, ext: str) -> bool:
        """Determine optimal processing method"""
        return ext == 'pdf'  # PDF is CPU-bound

    def _get_extension(self, path: str) -> str:
        """Get normalized file extension"""
        return removeFirstDot(getFileExtension(path)).lower()

    def shutdown(self):
        """Clean shutdown of processing system"""
        self.is_running = False
        if self.processing_thread and self.processing_thread.is_alive():
            self.processing_thread.join(timeout=5)
        if self.executor:
            self.executor.shutdown(wait=False)


# =====================
# Resource Monitor

class ResourceMonitor:
    """Monitors system resources to prevent overload"""

    def __init__(self, max_mem_usage=0.95, max_cpu_usage=0.8):
        self.max_mem = max_mem_usage
        self.max_cpu = max_cpu_usage
        self.last_check = 0
        self.last_result = False
        self.check_interval = 2  # seconds between checks

    def is_system_overloaded(self) -> bool:
        """Check if system resources are critically low"""
        now = time.time()
        if now - self.last_check < self.check_interval:
            return self.last_result

        try:
            mem = psutil.virtual_memory()
            cpu = psutil.cpu_percent(interval=0.5) / 100  # Longer interval for more accuracy

            self.last_result = (mem.percent / 100 > self.max_mem) or (cpu > self.max_cpu)
            self.last_check = now

            if self.last_result:
                print(f"System overloaded - Mem: {mem.percent}%, CPU: {cpu * 100:.1f}%")

            return self.last_result
        except Exception as e:
            print(f"Resource check failed: {e}")
            return False

# =====================
# Configuration
# =====================
class ThumbnailConfig:
    def __init__(self):
        # Processing parameters
        self.max_threads = 4
        self.batch_size = 10
        self.task_timeout = 120  # seconds

        # PDF settings
        self.pdf_zoom = 2.0

        # DOCX settings
        self.docx_zoom = 3.0
        self.docx_image_size = (800, 600)


# =====================
# Extractor Registry
# =====================
class ExtractorRegistry:
    _extractors: Dict[str, Type[DocExtractorABS]] = {}
    _lock = threading.Lock()

    @classmethod
    def register(cls, extensions: List[str]):
        """Decorator to register extractors"""

        def decorator(extractor_class: Type[DocExtractorABS]):
            with cls._lock:
                for ext in extensions:
                    cls._extractors[ext] = extractor_class
            return extractor_class

        return decorator

    @classmethod
    def get_extractor(cls, extension: str) -> Optional[Type[DocExtractorABS]]:
        """Get extractor class for file extension"""
        return cls._extractors.get(extension.lower())

    @classmethod
    def supported_extensions(cls) -> List[str]:
        """Get all supported extensions"""
        return list(cls._extractors.keys())


# =====================
# DocumentIconExtractor (Simplified Client)
# =====================
class DocumentIconExtractor(BaseGen):
    """Simplified client interface to the thumbnail system"""
    # Note this logic Set item_path to path if only trying to gen thumbnail_url

    def __init__(self, doc_path: str = '',
                 server_ip=NetworkConfig.server_ip, server_port=NetworkConfig.port):
        super().__init__(server_ip, server_port)
        self.task_queue = ThumbnailTaskQueue()

        self._extension=None
        self.document_path=doc_path
        if doc_path:
            self.add_document(doc_path)
            self.extension = self.document_path

    @property
    def item_path(self):
        return self.document_path

    def add_document(self, document_path: str):
        """Add document to processing queue"""
        self.document_path = document_path
        self.task_queue.add_task(document_path, self.thumbnail_path)

    def extract(self):
        """For API compatibility - processing happens automatically"""
        print('got here 2')

        pass  # Processing is managed by the queue system

    @property
    def pc_static_img(self):
        if self.extension not in self.task_queue.extractor_registry.supported_extensions():
            return "assets/icons/file.png"
        return use_pc_for_static(f"{self.extension}.png")

    @property
    def extension(self):
        return self._extension

    @extension.setter
    def extension(self, document_path):
        if document_path == '':
            self._extension = ''
            return
        self._extension = removeFirstDot(getFileExtension(document_path).lower())


# =====================
# Document Extractors
# =====================
@ExtractorRegistry.register(['pdf'])
class PdfExtractor(DocExtractorABS):
    def default_img_name(self):
        return 'pdf.png'
    def __init__(self):
        super().__init__()
        # Remove any thread locks or unpickleable attributes
        self._config = None  # Will be set in extract()

    def extract(self, absolute_filepath: str, thumbnail_path: str, config: ThumbnailConfig):
        try:
            # Make config available for this call
            self._config = config

            doc = fitz.open(absolute_filepath)
            page = doc.load_page(0)
            mat = fitz.Matrix(self._config.pdf_zoom, self._config.pdf_zoom)
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes()))
            img.save(thumbnail_path)
            return True
        except Exception as e:
            print(f"PDF extraction failed for file_path: {absolute_filepath} - {e}")
            traceback.print_exc()
            return self.getfailSafeImg()
        finally:
            self._config = None  # Clean up

@ExtractorRegistry.register(['docx', 'doc'])
class DocxExtractor(DocExtractorABS):
    def __init__(self):
        super().__init__()
        self._font_path = get_system_font()

    @property
    def default_img_name(self):
        return 'image.png'

    def extract(self, absolute_filepath: str, thumbnail_path: str, config: ThumbnailConfig):
        print('got here 1')
        try:
            doc = Document(absolute_filepath)
            img = Image.new('RGB', config.docx_image_size, (255, 255, 255))
            draw = ImageDraw.Draw(img)

            # Get first paragraph text
            text = doc.paragraphs[0].text if doc.paragraphs else "Word Document"

            # Configure font
            font_size = int(40 * config.docx_zoom)
            try:
                font = ImageFont.truetype(self._font_path, font_size) if self._font_path else ImageFont.load_default()
            except Exception as e:
                print(f"Custom PDF font paths not found: {self._font_path} - {e}")
                font = ImageFont.load_default()

            draw.text((50, 50), text, fill=(0, 0, 0), font=font)
            img.save(thumbnail_path)
            return True
        except Exception as e:
            print(f"DOCX extraction failed: {absolute_filepath} - {e}")
            return self.getfailSafeImg()


# =====================
# Initialization
# =====================
# Load any plugin extractors (implementation would be similar to before)
